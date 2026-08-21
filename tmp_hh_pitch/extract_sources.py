from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader


SOURCES = [
    Path("/Users/leiyi/Downloads/智能纪要：品牌150周年拍摄项目规划 2026年8月5日.pdf"),
    Path("/Users/leiyi/Downloads/Helly_Hansen_150周年拍摄项目策划方案.pdf"),
    Path("/Users/leiyi/Downloads/HH150周年《航海志》世界观深化创意文档_V2.docx"),
    Path("/Users/leiyi/Downloads/HH150周年《航海志》世界观深化创意文档.docx"),
    Path("/Users/leiyi/Downloads/HH150周年拍摄项目创意提案.docx"),
    Path("/Users/leiyi/Desktop/HH/HELLY HANSEN27SS代言人拍摄SHOOTING BRIEF (1).pdf"),
]

OUT = Path("/Users/leiyi/Documents/ChatGPT/HH拍摄pitch/tmp_hh_pitch/text")


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks = []
    for index, page in enumerate(reader.pages, start=1):
        chunks.append(f"\n\n===== PAGE {index} =====\n")
        chunks.append(page.extract_text() or "[NO EXTRACTABLE TEXT]")
    return "".join(chunks)


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    chunks = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            style_name = paragraph.style.name if paragraph.style is not None else "No Style"
            chunks.append(f"[{style_name}] {text}")
    for table_index, table in enumerate(doc.tables, start=1):
        chunks.append(f"\n===== TABLE {table_index} =====")
        for row in table.rows:
            chunks.append(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
    with ZipFile(path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    chunks.append(f"\n===== EMBEDDED MEDIA: {len(media)} =====")
    chunks.extend(media)
    return "\n".join(chunks)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for index, path in enumerate(SOURCES, start=1):
        if path.suffix.lower() == ".pdf":
            body = extract_pdf(path)
        else:
            body = extract_docx(path)
        target = OUT / f"{index:02d}_{path.stem}.txt"
        target.write_text(f"SOURCE: {path}\n{body}", encoding="utf-8")
        print(f"{target.name}\t{len(body)} chars")


if __name__ == "__main__":
    main()
